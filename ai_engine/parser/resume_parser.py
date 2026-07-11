"""
Resume Parser — extracts structured data from PDF/DOCX/image resumes
Uses PyMuPDF, python-docx, regex patterns + LLM for structured extraction
"""

import re
from pathlib import Path
from typing import Optional
from loguru import logger

from ai_engine.ocr.ocr_engine import OCREngine
from ai_engine.prompts.prompt_templates import RESUME_PARSE_PROMPT


class ResumeParser:
    """
    Multi-format resume parser supporting PDF, DOCX, and images.
    Uses rule-based extraction + LLM for structured output.
    """

    def __init__(self):
        self.ocr = OCREngine()

    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract raw text from any supported resume format."""
        ext = file_type.lower().lstrip(".")
        
        if ext == "pdf":
            return self._extract_pdf(file_path)
        elif ext in ["docx", "doc"]:
            return self._extract_docx(file_path)
        elif ext in ["png", "jpg", "jpeg", "tiff"]:
            return self.ocr.extract_from_image(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}, trying OCR")
            return self.ocr.extract_text(file_path)

    def _extract_pdf(self, file_path: str) -> str:
        try:
            import fitz  # lazy import PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                page_text = page.get_text("text")
                if len(page_text.strip()) < 30:
                    # Use OCR for scanned pages
                    text += self.ocr.extract_from_pdf(file_path)
                    break
                text += page_text + "\n"
            doc.close()
            return text.strip()
        except ImportError:
            logger.warning("PyMuPDF not installed, falling back to OCR")
            return self.ocr.extract_from_pdf(file_path)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return self.ocr.extract_from_pdf(file_path)

    def _extract_docx(self, file_path: str) -> str:
        try:
            from docx import Document  # lazy import python-docx
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx not installed")
            return ""
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""

    def parse_with_rules(self, text: str) -> dict:
        """Rule-based parsing for contact information."""
        return {
            "name": self._extract_name(text),
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "location": self._extract_location(text),
            "linkedin": self._extract_linkedin(text),
            "github": self._extract_github(text),
        }

    def _extract_email(self, text: str) -> Optional[str]:
        pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
        match = re.search(pattern, text)
        return match.group(0) if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        pattern = r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"
        match = re.search(pattern, text)
        return match.group(0) if match else None

    def _extract_name(self, text: str) -> Optional[str]:
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:5]:
            # Skip lines that look like titles/sections
            if len(line.split()) in (2, 3) and not any(
                kw in line.lower() for kw in ["resume", "cv", "curriculum", "address", "@"]
            ):
                if re.match(r"^[A-Za-z\s\.\-']+$", line):
                    return line
        return lines[0] if lines else None

    def _extract_location(self, text: str) -> Optional[str]:
        pattern = r"[A-Z][a-zA-Z\s]+,\s*[A-Z]{2}(?:\s+\d{5})?"
        match = re.search(pattern, text)
        return match.group(0) if match else None

    def _extract_linkedin(self, text: str) -> Optional[str]:
        pattern = r"linkedin\.com/in/[\w\-]+"
        match = re.search(pattern, text, re.IGNORECASE)
        return match.group(0) if match else None

    def _extract_github(self, text: str) -> Optional[str]:
        pattern = r"github\.com/[\w\-]+"
        match = re.search(pattern, text, re.IGNORECASE)
        return match.group(0) if match else None

    def parse_with_llm(self, text: str, llm) -> dict:
        """Use LLM to extract structured resume sections."""
        prompt = RESUME_PARSE_PROMPT.format(resume_text=text[:3000])  # Token limit
        response = llm.invoke(prompt)
        
        try:
            import json
            # Try to extract JSON from response
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
        except Exception:
            pass

        # Fallback to rule-based extraction
        return self._fallback_section_parse(text)

    def _fallback_section_parse(self, text: str) -> dict:
        """Simple section-based parsing as fallback."""
        sections = {
            "summary": "",
            "education": [],
            "experience": [],
            "projects": [],
            "skills": [],
            "certifications": [],
        }

        # Skills extraction by common keywords
        skill_keywords = [
            "Python", "Java", "JavaScript", "TypeScript", "React", "Node.js",
            "SQL", "PostgreSQL", "MongoDB", "Docker", "Kubernetes", "AWS",
            "Azure", "GCP", "FastAPI", "Django", "Flask", "Machine Learning",
            "TensorFlow", "PyTorch", "LangChain", "OpenAI", "Git", "REST API",
            "GraphQL", "Redis", "Elasticsearch", "Pandas", "NumPy", "Spark",
            "Scala", "Go", "Rust", "C++", "C#", ".NET", "Spring Boot",
        ]
        found_skills = [kw for kw in skill_keywords if kw.lower() in text.lower()]
        sections["skills"] = found_skills

        return sections

    async def parse(self, file_path: str, file_type: str, llm=None) -> dict:
        """Full async parse: extract text + structured data."""
        raw_text = self.extract_text(file_path, file_type)
        
        contact_info = self.parse_with_rules(raw_text)
        
        if llm:
            structured = self.parse_with_llm(raw_text, llm)
        else:
            structured = self._fallback_section_parse(raw_text)

        result = {
            **contact_info,
            **structured,
            "raw_text": raw_text,
        }
        
        logger.info(f"✅ Resume parsed: {result.get('name', 'Unknown')}, {len(result.get('skills', []))} skills")
        return result
